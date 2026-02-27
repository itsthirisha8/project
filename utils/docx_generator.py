from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json

def generate_resume_docx(resume_data):
    doc = Document()
    
    # Set standard font for ATS
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    content = json.loads(resume_data.content)
    personal = content.get('personal', {})
    
    # Header
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(personal.get('full_name', ''))
    run.bold = True
    run.font.size = Pt(20)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f"{personal.get('location', '')} | {personal.get('phone', '')} | {personal.get('email', '')}"
    if personal.get('linkedin'):
        contact_text += f" | {personal.get('linkedin', '')}"
    contact.add_run(contact_text)

    # Summary
    if content.get('summary'):
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(content.get('summary'))

    # Experience
    if content.get('experience'):
        doc.add_heading('Experience', level=1)
        for exp in content.get('experience'):
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('title')} - {exp.get('company')}")
            run.bold = True
            p.add_run(f"\t{exp.get('dates')}")
            
            sub = doc.add_paragraph()
            sub.add_run(f"{exp.get('location')}").italic = True
            
            doc.add_paragraph(exp.get('description'))

    # Education
    if content.get('education'):
        doc.add_heading('Education', level=1)
        for edu in content.get('education'):
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('degree')} - {edu.get('school')}")
            run.bold = True
            p.add_run(f"\t{edu.get('dates')}")
            doc.add_paragraph(edu.get('location')).italic = True

    # Projects
    if content.get('projects'):
        doc.add_heading('Projects', level=1)
        for prj in content.get('projects'):
            p = doc.add_paragraph()
            p.add_run(prj.get('name')).bold = True
            doc.add_paragraph(prj.get('description'))

    # Skills
    if content.get('skills'):
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(content.get('skills'))

    # Certifications
    if content.get('certifications'):
        doc.add_heading('Certifications', level=1)
        for cert in content.get('certifications'):
            doc.add_paragraph(f"{cert.get('title')} - {cert.get('issuer')}")

    # Achievements
    if content.get('achievements'):
        doc.add_heading('Achievements', level=1)
        doc.add_paragraph(content.get('achievements'))

    # Languages
    if content.get('languages'):
        doc.add_heading('Languages', level=1)
        doc.add_paragraph(content.get('languages'))

    # Save to buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
